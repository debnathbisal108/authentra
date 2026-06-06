import io
import tempfile
import os
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_data: bytes) -> str:
    """Extract text from PDF using pdfplumber with PyMuPDF fallback and OCR for scanned pages."""
    text_parts = []
    
    # Try pdfplumber first
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                else:
                    # Page might be scanned - try OCR
                    ocr_text = _ocr_page_from_pdf(file_data, page.page_number - 1)
                    if ocr_text:
                        text_parts.append(ocr_text)
        
        if text_parts:
            return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Fallback to PyMuPDF
    try:
        import fitz
        doc = fitz.open(stream=file_data, filetype="pdf")
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
            else:
                # Try OCR via PyMuPDF pixmap
                try:
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    ocr_text = _ocr_image_bytes(img_bytes)
                    if ocr_text:
                        text_parts.append(ocr_text)
                except Exception as ocr_e:
                    logger.warning(f"OCR failed for page {page_num}: {ocr_e}")
        doc.close()
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PyMuPDF failed: {e}")
        return ""


def _ocr_page_from_pdf(file_data: bytes, page_num: int) -> str:
    """OCR a specific page from a PDF."""
    try:
        import fitz
        doc = fitz.open(stream=file_data, filetype="pdf")
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        doc.close()
        return _ocr_image_bytes(img_bytes)
    except Exception as e:
        logger.warning(f"OCR page extraction failed: {e}")
        return ""


def _ocr_image_bytes(img_bytes: bytes) -> str:
    """Run OCR on image bytes using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        return pytesseract.image_to_string(img, lang="eng")
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text(file_data: bytes, file_type: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from resume file.
    Returns (extracted_text, error_message)
    """
    file_type_lower = file_type.lower().strip(".")
    
    if file_type_lower == "pdf":
        text = extract_text_from_pdf(file_data)
    elif file_type_lower in ("docx", "doc"):
        text = extract_text_from_docx(file_data)
    else:
        return "", f"Unsupported file type: {file_type}"
    
    if not text or not text.strip():
        return "", "Could not extract text from file. The file may be corrupted or in an unsupported format."
    
    return text.strip(), None
